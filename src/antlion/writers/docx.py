from pathlib import Path

from docx import Document


def write_docx(content: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    doc.save(str(path))
